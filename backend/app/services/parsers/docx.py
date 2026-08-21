"""Deterministic DOCX parser using python-docx."""

import io
import zipfile

import docx

from app.services.cleaner import (
    dehyphenate_text,
    detect_language,
    normalize_whitespace,
)
from app.services.parsers.base import BaseDocumentParser, ParsedDocument, ParsedSection

# OLE2 Compound Document binary header for legacy .doc
OLE2_DOC_MAGIC = b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1"


class DOCXDocumentParser(BaseDocumentParser):
    """Parser extracting structured text, headings, and tables from DOCX documents."""

    def parse(self, content_bytes: bytes, filename: str) -> ParsedDocument:
        # Detect and reject legacy .doc binary files immediately
        if content_bytes.startswith(OLE2_DOC_MAGIC) or filename.lower().endswith(".doc"):
            return ParsedDocument(
                error_code="UNSUPPORTED_LEGACY_FORMAT",
                error_message=(
                    "Legacy .doc format is not supported. "
                    "Please convert your file to .docx or .pdf before uploading."
                ),
            )

        # Protect against zip bombs / corrupted archives
        try:
            with zipfile.ZipFile(io.BytesIO(content_bytes)) as zf:
                total_uncompressed = sum(info.file_size for info in zf.infolist())
                if total_uncompressed > 200 * 1024 * 1024:  # 200MB expansion limit
                    return ParsedDocument(
                        error_code="ZIP_BOMB_DETECTED",
                        error_message="Document exceeds uncompressed security size limits.",
                    )
        except Exception as e:
            return ParsedDocument(
                error_code="CORRUPTED_FILE",
                error_message=f"Failed to read DOCX archive: {str(e)}",
            )

        try:
            doc = docx.Document(io.BytesIO(content_bytes))
        except Exception as e:
            return ParsedDocument(
                error_code="CORRUPTED_FILE",
                error_message=f"Failed to parse DOCX document: {str(e)}",
            )

        sections: list[ParsedSection] = []
        current_title: str = "Document Overview"
        current_paragraphs: list[str] = []
        current_level: int = 1

        for para in doc.paragraphs:
            text = normalize_whitespace(dehyphenate_text(para.text))
            if not text:
                continue

            style_name = para.style.name if para.style else ""
            if "Heading" in style_name or "Title" in style_name:
                if current_paragraphs:
                    section_content = "\n\n".join(current_paragraphs)
                    sections.append(
                        ParsedSection(
                            title=current_title,
                            content=section_content,
                            page_start=1,
                            page_end=1,
                            level=current_level,
                        )
                    )
                    current_paragraphs = []

                current_title = text
                if "1" in style_name:
                    current_level = 1
                elif "2" in style_name:
                    current_level = 2
                elif "3" in style_name:
                    current_level = 3
                else:
                    current_level = 1
            else:
                current_paragraphs.append(text)

        # Extract tables
        for table in doc.tables:
            table_rows: list[str] = []
            for row in table.rows:
                cells = [
                    normalize_whitespace(cell.text)
                    for cell in row.cells
                    if normalize_whitespace(cell.text)
                ]
                if cells:
                    table_rows.append(" | ".join(cells))
            if table_rows:
                current_paragraphs.append("Table Data:\n" + "\n".join(table_rows))

        if current_paragraphs:
            section_content = "\n\n".join(current_paragraphs)
            sections.append(
                ParsedSection(
                    title=current_title,
                    content=section_content,
                    page_start=1,
                    page_end=1,
                    level=current_level,
                )
            )

        full_raw_text = "\n\n".join(f"{s.title}\n{s.content}" for s in sections)
        word_count = len(full_raw_text.split())
        language = detect_language(full_raw_text)

        # Estimate page count (roughly 350 words per standard page)
        estimated_pages = max(1, (word_count // 350) + 1)

        return ParsedDocument(
            sections=sections,
            page_count=estimated_pages,
            word_count=word_count,
            language=language,
            raw_text=full_raw_text,
            is_scanned=False,
            metadata={"parser": "python-docx"},
        )
